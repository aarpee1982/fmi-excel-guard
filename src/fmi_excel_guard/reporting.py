from __future__ import annotations

from io import BytesIO

import pandas as pd
from docx import Document

from .models import Finding


def findings_to_dataframe(findings: list[Finding]) -> pd.DataFrame:
    rows = [
        {
            "Market Name": finding.market_name,
            "Category": finding.category,
            "Find": finding.find_text,
            "Replace with": finding.replace_with,
            "Why flagged": finding.why_flagged,
            "Source": finding.source,
            "Confidence": round(finding.confidence, 2),
        }
        for finding in findings
    ]
    return pd.DataFrame(rows)


def build_findings_docx(findings: list[Finding]) -> bytes:
    document = Document()
    document.add_heading("FMI Upload Guard: Issues and Suggested Fixes", level=1)
    document.add_paragraph(f"Total findings: {len(findings)}")

    grouped: dict[str, list[Finding]] = {}
    for finding in findings:
        grouped.setdefault(finding.market_name, []).append(finding)

    for market_name in sorted(grouped):
        document.add_heading(market_name, level=2)
        for index, finding in enumerate(grouped[market_name], start=1):
            document.add_heading(f"Issue {index}", level=3)
            document.add_paragraph(f"Find: {finding.find_text}")
            document.add_paragraph(f"Replace with: {finding.replace_with}")
            document.add_paragraph(f"Why flagged: {finding.why_flagged}")
            document.add_paragraph(f"Source: {finding.source} ({finding.confidence:.2f})")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
