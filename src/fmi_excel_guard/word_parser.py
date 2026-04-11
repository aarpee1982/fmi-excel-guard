from __future__ import annotations

import re
from io import BytesIO
from typing import Iterable

from docx import Document

from .models import FAQItem, MarketRecord, Section
from .parser import normalize_text


def load_market_records_from_word_files(files: Iterable[tuple[str, bytes]]) -> list[MarketRecord]:
    records: list[MarketRecord] = []
    for index, (filename, payload) in enumerate(files, start=1):
        document = Document(BytesIO(payload))
        records.append(_build_record_from_document(filename=filename, document=document, index=index))
    return records


def load_market_record_from_text(*, text: str, title: str = "Pasted Article") -> MarketRecord:
    paragraphs = [normalize_text(part) for part in text.splitlines()]
    paragraphs = [paragraph for paragraph in paragraphs if paragraph]
    return _build_record_from_paragraphs(filename=title, paragraphs=paragraphs, sections=_sections_from_paragraphs(paragraphs), index=1)


def _build_record_from_document(*, filename: str, document: Document, index: int) -> MarketRecord:
    paragraphs = [normalize_text(paragraph.text) for paragraph in document.paragraphs]
    paragraphs = [paragraph for paragraph in paragraphs if paragraph]

    heading_sections: list[Section] = []
    current_title = "Document Body"
    current_lines: list[str] = []

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if style_name.startswith("heading"):
            if current_lines:
                heading_sections.append(Section(title=current_title, text=" ".join(current_lines)))
            current_title = text
            current_lines = []
            continue
        current_lines.append(text)

    if current_lines:
        heading_sections.append(Section(title=current_title, text=" ".join(current_lines)))

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            values = [normalize_text(cell.text) for cell in row.cells]
            values = [value for value in values if value]
            if values:
                rows.append(" | ".join(values))
        if rows:
            heading_sections.append(Section(title=f"Table {table_index}", text=" ".join(rows)))

    return _build_record_from_paragraphs(
        filename=filename,
        paragraphs=paragraphs,
        sections=heading_sections,
        index=index,
    )


def _build_record_from_paragraphs(
    *,
    filename: str,
    paragraphs: list[str],
    sections: list[Section],
    index: int,
) -> MarketRecord:
    market_name = _extract_market_name(filename, paragraphs, sections)
    meta_title = paragraphs[0] if paragraphs else market_name
    meta_desc = _extract_primary_summary(paragraphs)
    rep_sub_title = _extract_secondary_summary(paragraphs)
    rep_title = " ".join(paragraphs[:8])
    toc_text = " ".join(section.title for section in sections)
    faq_items = _extract_faq_items(paragraphs, sections)

    return MarketRecord(
        rep_id=index,
        market_name=market_name,
        meta_desc=meta_desc,
        meta_title=meta_title,
        rep_title=rep_title,
        rep_sub_title=rep_sub_title,
        toc_text=toc_text,
        faq_items=faq_items,
        description_sections=sections,
    )


def _sections_from_paragraphs(paragraphs: list[str]) -> list[Section]:
    if not paragraphs:
        return []
    return [Section(title="Document Body", text=" ".join(paragraphs))]


def _extract_market_name(filename: str, paragraphs: list[str], sections: list[Section]) -> str:
    for text in [*(section.title for section in sections[:6]), *paragraphs[:6]]:
        match = re.search(r"([A-Z][A-Za-z0-9/&(),.'\- ]+ Market)", text)
        if match:
            return normalize_text(match.group(1))
    return normalize_text(filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " "))


def _extract_primary_summary(paragraphs: list[str]) -> str:
    for paragraph in paragraphs:
        lowered = paragraph.lower()
        if "usd" in lowered and ("cagr" in lowered or "projected to reach" in lowered or "forecast" in lowered):
            return paragraph
    return paragraphs[0] if paragraphs else ""


def _extract_secondary_summary(paragraphs: list[str]) -> str:
    candidates = []
    for paragraph in paragraphs:
        lowered = paragraph.lower()
        if "usd" in lowered or "cagr" in lowered or "forecast period" in lowered:
            candidates.append(paragraph)
    return " ".join(candidates[:2]) if candidates else (paragraphs[1] if len(paragraphs) > 1 else "")


def _extract_faq_items(paragraphs: list[str], sections: list[Section]) -> list[FAQItem]:
    faq_items: list[FAQItem] = []
    section_text = " ".join(section.title.lower() for section in sections)
    if "faq" not in section_text and "frequently asked" not in section_text:
        return faq_items

    for index, paragraph in enumerate(paragraphs[:-1]):
        if paragraph.endswith("?"):
            answer = paragraphs[index + 1]
            faq_items.append(FAQItem(question=paragraph, answer=answer))
    return faq_items[:12]
